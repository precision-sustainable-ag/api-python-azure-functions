import json
import os
import logging
import os
import traceback
import pandas as pd
from docx import Document
from docx.shared import Inches
import azure.functions as func
from datetime import datetime, timedelta
# import matplotlib.pyplot as plt 

from SharedFunctions import db_connectors, global_vars, initializer
from FetchReport.utils import biomass, moisture, precipitation, crop_yield, hyperlink, gdd

class FetchReport:
    def __init__(self, req):
        initial_state = initializer.initialize(
            route_params=["site"], body_params=None, req=req)

        self.route_params_obj = initial_state["route_params_obj"]
        self.token = initial_state["token"]
        self.token_missing = initial_state["token"] == None
        self.authenticated = initial_state["authenticated"]
        self.auth_response = initial_state["auth_response"]
        self.invalid_params = initial_state["route_params_obj"] is None

        # connect to dbs
        if self.authenticated:
            self.environment = os.environ.get('AZURE_FUNCTIONS_ENVIRONMENT')
            self.shadow_con, self.shadow_cur, self.shadow_engine = db_connectors.connect_to_crown(
                self.environment)

        self.requested_site = self.route_params_obj.get("site")

    def fetch_reportdata(self):
        report_data = pd.DataFrame(pd.read_sql(
            "SELECT a.code, a.affiliation, a.county, a.longitude, a.latitude, "+
            "a.address, b.cc_planting_date, b.cc_termination_date, "+
            "b.cash_crop_planting_date, b.cash_crop_harvest_date FROM "+
            "site_information as a INNER JOIN farm_history as b ON a.code=b.code"+
            " WHERE a.code = '{}'"\
                .format(self.route_params_obj.get("site")), self.shadow_engine))
        return report_data

    def generate_report(self):
        report_data = self.fetch_reportdata()

        cash_planting = report_data.iloc[0].get("cash_crop_planting_date") \
            if report_data.iloc[0].get("cash_crop_planting_date") != None else None
        cash_harvest = report_data.iloc[0].get("cash_crop_harvest_date") \
            if report_data.iloc[0].get("cash_crop_harvest_date") != None else None
        cover_planting = report_data.iloc[0].get("cc_planting_date") \
            if report_data.iloc[0].get("cc_planting_date") != None else None
        cover_termination = report_data.iloc[0].get("cc_termination_date") \
            if report_data.iloc[0].get("cc_termination_date") != None else None

        lat = round(report_data.iloc[0].get("latitude"), 4)
        lon = round(report_data.iloc[0].get("longitude"), 4)

        site_biomass, species = biomass.fetch_biomass(report_data.iloc[0].get("affiliation"), self.requested_site)
        bare_yield, cover_yield = crop_yield.fetch_yield(self.requested_site)
        cash_precipitation = precipitation.fetch_precipitation(cash_planting, cash_harvest, lat, lon) \
            if (cash_planting!=None and cash_harvest!=None) else None
        cover_precipitation = precipitation.fetch_precipitation(cover_planting, cover_termination, lat, lon) \
            if (cover_planting!=None and cover_termination!=None) else None

        cash_gdd = gdd.fetch_gdd(cash_planting, cash_harvest, lat, lon, 10) \
            if (cash_planting!=None and cash_harvest!=None) else None 
        cover_gdd = gdd.fetch_gdd(cover_planting, cover_termination, lat, lon, 4) \
            if (cover_planting!=None and cover_termination!=None) else None

        cash_days = (cash_harvest-cash_planting).days \
            if (cash_planting!=None and cash_harvest!=None) else None 
        cover_days = (cover_termination-cover_planting).days \
            if (cover_planting!=None and cover_termination!=None) else None

        vwc = moisture.fetch_vwc(cover_planting, cover_termination, self.requested_site)\
            if (cover_planting!=None and cover_termination!=None) else None
        vwc_dict = {}
        if vwc.size != 0:
            vwc['date'] = pd.to_datetime(vwc['timestamp'])
            newdf = (vwc.groupby(['treatment', pd.Grouper(key='date', freq='W')]).agg({'vwc':'mean'}))
            newdf = newdf.reset_index()
            newdf['date'] = pd.to_datetime(newdf['date']).dt.date
            newdf['date'] = pd.to_datetime(newdf['date']).dt.strftime('%m/%d/%Y')
            print(newdf)

            for i in range(len(newdf)):
                d = str(newdf.iloc[i].get('date'))
                # print(d)
                if d not in vwc_dict.keys():
                    vwc_dict[d] = ["", ""]
                if str(newdf.iloc[i].get('treatment')) == "b":
                    vwc_dict[d][0] = str(round(newdf.iloc[i].get('vwc'), 3))
                elif str(newdf.iloc[i].get('treatment')) == "c":
                    vwc_dict[d][1] = str(round(newdf.iloc[i].get('vwc'), 3))

        if report_data.empty:
            return func.HttpResponse(json.dumps({"status": "error", "details": "no site code in crown db"}), headers=global_vars.HEADER, status_code=400)
        else:
            doc = Document()

            #header section
            #adding header with PSA logo
            section = doc.sections[0]
            header = section.header
            headerPara = header.paragraphs[0]
            headerLogo = headerPara.add_run()
            headerLogo.add_picture("FetchReport\\PSA.png", width=Inches(1))

            #footer
            footer = section.footer
            footerPara = footer.paragraphs[0]
            footerPara.add_run('For any grievances reaach us at : abc@xyz.com')

            #PSA paragraph
            psaPara = doc.add_paragraph('\n'+'The Precision Sustainable Agriculture (PSA)'+ 
            'On-Farm network deploys common research protocols that study the short-term'+
            ' effects of cover crops on farms that currently use cover crops. '+
            'By utilizing farms with different management practices, the data collected'+
            ' can account for a wide range of factors such as termination timing, specific selections, and climate impacts. '+
            'The data is used to build tools to aid in site-specific management decisions.\n')

            #Farm Name
            doc.add_heading('Farm Name:', 3)
            farmName = doc.add_paragraph().add_run(report_data.iloc[0].get("code"))

            #Farm Address
            doc.add_heading('Farm Address:', 3)
            farmAdd = doc.add_paragraph().add_run(report_data.iloc[0].get("address"))

            #Site Description
            doc.add_heading('Site Description:', 3)
            gpsPara = doc.add_paragraph()
            gpsPara.add_run('GPS Co-ordinates\n').itallic=True
            gpsPara.add_run('Latitude: ')
            gpsPara.add_run(str(lat))
            gpsPara.add_run('\t')
            gpsPara.add_run('Longitude: ')
            gpsPara.add_run(str(lon))
            gpsPara.add_run('\n')

            maps_link = "https://www.google.com/maps/search/{lat},{lon}"\
                .format(lat=lat, lon=lon)
            hyperlink.add_hyperlink(gpsPara, maps_link, maps_link)

            #Dates this report summarizes
            doc.add_heading('Dates this report summarizes:', 3)
            currentDate = datetime.now()
            date = doc.add_paragraph(currentDate.strftime("%m/%d/%Y"))

            #Summary of Activities and Management
            doc.add_heading('Summary of Activities and Management:', 3)
            # activitesPara = doc.add_paragraph()

            # Cash crop dates
            doc.add_paragraph('Date of planting Cash crop: ', style='List Bullet'
            ).add_run(cash_planting.strftime("%m/%d/%Y") if cash_planting else "Not yet entered")
            doc.add_paragraph('Date of harvest Cash crop: ', style='List Bullet'
            ).add_run(cash_harvest.strftime("%m/%d/%Y") if cash_harvest else "Not yet entered")
            doc.add_paragraph('Cash crop no of days in production: ', style='List Bullet'
            ).add_run(str(cash_days) if cash_days else "________")

            #Cover crop dates
            doc.add_paragraph('Date of planting Cover crop: ', style='List Bullet'
            ).add_run(cover_planting.strftime("%m/%d/%Y") if cover_planting else "Not yet entered")
            doc.add_paragraph('Date of termination Cover crop: ', style='List Bullet'
            ).add_run(cover_termination.strftime("%m/%d/%Y") if cover_termination else "Not yet entered")
            doc.add_paragraph('Cover crop no of days in production: ', style='List Bullet'
            ).add_run(str(cover_days) if cover_days else "________")

            #GDD
            gdd_para = doc.add_paragraph('Total GDD: ', style='List Bullet')
            gdd_para.add_run("\nGDD for Cover crop is ")
            gdd_para.add_run(str(round(cover_gdd[0]['sum(gdd)'], 1)) if cover_gdd else "________")
            gdd_para.add_run("\nGDD for Cash crop is ")
            gdd_para.add_run(str(round(cash_gdd[0]['sum(gdd)'], 1)) if cash_gdd else "________")

            #Precipitation
            preci_para = doc.add_paragraph('Total precipitation: ', style='List Bullet')
            preci_para.add_run("\nPrecipitation for Cover crop is ")
            preci_para.add_run(str(round(cover_precipitation[0]['sum(precipitation)'], 1))+" mm" if cover_precipitation else "________")
            preci_para.add_run("\nPrecipitation for Cash crop is ")
            preci_para.add_run(str(round(cash_precipitation[0]['sum(precipitation)'], 1))+" mm" if cash_precipitation else "________")

            #Biomass and comparison
            species = species if species else "Unavailable"
            doc.add_paragraph('Cover crop species: ', style='List Bullet'
            ).add_run(species)
            
            site_biomass = str(round(site_biomass, 1)) if site_biomass else "Unavailable"
            doc.add_paragraph('Dry matter (lbs/acre):', style='List Bullet'
            ).add_run(site_biomass)
            biomass_comp_para = doc.add_paragraph('Dry matter in comparison to others in the region:'+
            ' \n', style='List Bullet')
            if os.path.exists("FetchReport\\Graph.png"):
                doc.add_picture("FetchReport\\Graph.png")
            else:
                biomass_comp_para.add_run("Comparison not available")

            #Yield both bare ground and cover crop
            yield_para = doc.add_paragraph('Yield: ', style='List Bullet')
            yield_para.add_run("\nBare Ground :")
            yield_para.add_run(str(bare_yield)+" Mg/ha" if bare_yield else "Not available")
            yield_para.add_run("\nCover :")
            yield_para.add_run(str(bare_yield)+" Mg/ha" if cover_yield else "Not available")

            #Water and Moisture
            doc.add_heading('Soil Temperature: ', 3)
            doc.add_paragraph('Temperature data: ', style='List Bullet')
            if os.path.exists("FetchReport\\TemperatureGraph.png"):
                doc.add_picture("FetchReport\\TemperatureGraph.png")            
            #Water and Moisture
            doc.add_heading('Water and Moisture: ', 3)
            doc.add_paragraph('Moisture data: ', style='List Bullet')
            doc.add_paragraph('Cover crop vs bare: ', style='List Number 2').add_run("Refer the table below")
            table = doc.add_table(rows=1, cols=3, style="Table Grid")
            row = table.rows[0].cells
            row[0].text = 'Week'
            row[1].text = 'Bare Ground'
            row[2].text = 'Cover Crop'

            for key, value in vwc_dict.items():
                # Adding a row and then adding data in it.
                row = table.add_row().cells
                # Converting id to string as table can only take string input
                row[0].text = key
                row[1].text = value[0]
                row[2].text = value[1]

            if os.path.exists("FetchReport\\MoistureGraph.png"):
                doc.add_picture("FetchReport\\MoistureGraph.png")
            if os.path.exists("FetchReport\\MoistureGraphD.png"):
                doc.add_picture("FetchReport\\MoistureGraphD.png")
            if os.path.exists("FetchReport\\MoistureGraphC.png"):
                doc.add_picture("FetchReport\\MoistureGraphC.png")
            if os.path.exists("FetchReport\\MoistureGraphB.png"):
                doc.add_picture("FetchReport\\MoistureGraphB.png")
            if os.path.exists("FetchReport\\MoistureGraphA.png"):
                doc.add_picture("FetchReport\\MoistureGraphA.png")

            #Decision Support Tools
            doc.add_heading('Decision Support Tools:', 3)
            dstPara = doc.add_paragraph()
            dstPara.add_run('The Decision Support Tools (DSTs) are designed for farmers'+
            ' to input their data and receive custom generated information on how to '+
            'address their management strategies. The Cover Crop Nitrogen Calculator '+
            '(CC-NCALC) calculates the amount of nitrogen available after the planting '+
            'and termination of a cover crop. ')
            
            # now save the document to a location
            file_name = 'SummaryReport'+self.requested_site+'.docx'
            doc.save(file_name)
            if os.path.exists("FetchReport\\Graph.png"):
                os.remove("FetchReport\\Graph.png")
            if os.path.exists("FetchReport\\MoistureGraph.png"):
                os.remove("FetchReport\\MoistureGraph.png")
            return_obj = {
                "status": "success",
                "details": "successfully fetched report data for {}".format(self.route_params_obj.get("site")),
                "json_response": "response obtained from fetch query {}".format(report_data),
            }

            return func.HttpResponse(json.dumps(return_obj), headers=global_vars.HEADER, status_code=201)


def main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Python HTTP trigger function processed a request.')

    try:
        report_data_fetcher = FetchReport(req)

        # authenticate user based on token
        if any([not report_data_fetcher.authenticated, report_data_fetcher.invalid_params, report_data_fetcher.token_missing]):
            return initializer.get_response(report_data_fetcher.authenticated, report_data_fetcher.auth_response,
                                            report_data_fetcher.invalid_params, report_data_fetcher.token_missing, False)
        else:
            return report_data_fetcher.generate_report()

    except Exception as error:
        logging.info("program encountered exception: " +
                     traceback.format_exc())
        logging.exception(error)
        return func.HttpResponse(
            json.dumps({"status": "error", "details": traceback.format_exc()}),
            status_code=500,
            headers=global_vars.HEADER
        )
