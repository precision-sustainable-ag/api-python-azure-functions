from ..controller import raw_nir, crop_yield, moisture
from ..controller import gdd, precipitation, hyperlink, biomass
import os
# from docx import Document
from docx.shared import Inches
import pandas as pd
import numpy as np


def doc_header(doc):
    try:
        # header section
        # adding header with PSA logo
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_logo = header_para.add_run()
        header_logo.add_picture("FetchReport/PSA.png", width=Inches(1))

        # footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.add_run(
            'For questions about this report, ask: abc@xyz.com')
    except Exception as e:
        print(e)
    finally:
        return doc


def doc_farmDetails(doc, site_info):
    try:
        # Farm Name
        doc.add_heading('Farm Name:', 4)
        doc.add_paragraph().add_run(
            site_info.iloc[0].get("code"))

        # Farm Address
        doc.add_heading('Farm Address:', 4)
        doc.add_paragraph().add_run(
            site_info.iloc[0].get("address"))

        lat = round(site_info.iloc[0].get("latitude"), 4)
        lon = round(site_info.iloc[0].get("longitude"), 4)

        # Site Description
        doc.add_heading('Site Description:', 4)
        gps_para = doc.add_paragraph()
        gps_para.add_run('GPS Co-ordinates\n').itallic = True
        gps_para.add_run('Latitude: ')
        gps_para.add_run(str(lat))
        gps_para.add_run('\t')
        gps_para.add_run('Longitude: ')
        gps_para.add_run(str(lon))
        gps_para.add_run('\n')

        maps_link = "https://www.google.com/maps/search/{lat},{lon}"\
            .format(lat=lat, lon=lon)
        hyperlink.add_hyperlink(gps_para, maps_link, maps_link)
    except Exception as e:
        print(e)
    finally:
        return doc


def doc_cashcrop(doc, cash_planting, cash_harvest):
    try:
        cash_days = (cash_harvest-cash_planting).days if (cash_planting is not None and cash_harvest is not None) else None
        
        doc.add_heading('Cash Crop:', 4)
        # Cash crop dates
        doc.add_paragraph('Date of Planting Cash Crop: ', style='List Bullet'
                          ).add_run(cash_planting.strftime("%m/%d/%Y") if cash_planting else "Not yet entered")
        doc.add_paragraph('Date of Harvest Cash Crop: ', style='List Bullet'
                          ).add_run(cash_harvest.strftime("%m/%d/%Y") if cash_harvest else "Not yet entered")
        doc.add_paragraph('Number of Days in Production: ', style='List Bullet'
                          ).add_run(str(cash_days) if cash_days else "________")
        
    except Exception as e:
        print(e)
    finally:
        return doc


def doc_covercrop(doc, cover_planting, cover_termination):
    try:
        cover_days = (cover_termination-cover_planting).days \
            if (cover_planting is not None and cover_termination is not None) else None
        doc.add_heading('Cover Crop Days:', 4)
        # Cover crop dates
        doc.add_paragraph('Date of Olanting Cover Crop: ', style='List Bullet'
                          ).add_run(cover_planting.strftime("%m/%d/%Y") if cover_planting else "Not yet entered")
        doc.add_paragraph('Date of Termination Cover Crop: ', style='List Bullet'
                          ).add_run(cover_termination.strftime("%m/%d/%Y") if cover_termination else "Not yet entered")
        doc.add_paragraph('Number of Days in Production: ', style='List Bullet'
                          ).add_run(str(cover_days) if cover_days else "________")
    except Exception as e:
        print(e)
    finally:
        return doc


def doc_gdd(doc, cash_planting, cash_harvest, cover_planting, cover_termination, lat, lon):
    try:
        cash_gdd = gdd.fetch_gdd(cash_planting, cash_harvest, lat, lon, 10) \
            if (cash_planting is not None and cash_harvest is not None) else None
        cover_gdd = gdd.fetch_gdd(cover_planting, cover_termination, lat, lon, 4) \
            if (cover_planting is not None and cover_termination is not None) else None

        doc.add_heading('GDD:', 4)
        gdd_para = doc.add_paragraph(
            'GDD for Cover Crop is ', style='List Bullet')
        gdd_para.add_run(
            str(round(cover_gdd[0]['sum(gdd)'], 1)) if cover_gdd else "________")
        gdd_para = doc.add_paragraph(
            'GDD for Cash Crop is ', style='List Bullet')
        gdd_para.add_run(
            str(round(cash_gdd[0]['sum(gdd)'], 1)) if cash_gdd else "________")

    except Exception as e:
        print(e)
    finally:
        return doc


def doc_precipitation(doc, cash_planting, cash_harvest, cover_planting, cover_termination, lat, lon):
    try:
        cash_precipitation = precipitation.fetch_precipitation(cash_planting, cash_harvest, lat, lon) \
            if (cash_planting is not None and cash_harvest is not None) else None
        cover_precipitation = precipitation.fetch_precipitation(cover_planting, cover_termination, lat, lon) \
            if (cover_planting is not None and cover_termination is not None) else None

        doc.add_heading('Precipitation:', 4)
        # Precipitation
        preci_para = doc.add_paragraph(
            'Precipitation for Cover Crop is ', style='List Bullet')
        preci_para.add_run(str(round(
            cover_precipitation[0]['sum(precipitation)'], 1))+" mm" if cover_precipitation else "________")
        preci_para = doc.add_paragraph(
            'Precipitation for Cash Crop is ', style='List Bullet')
        preci_para.add_run(str(round(
            cash_precipitation[0]['sum(precipitation)'], 1))+" mm" if cash_precipitation else "________")

    except Exception as e:
        print(e)
    finally:
        return doc


def doc_biomass(doc, affiliation, year, requested_site):
    try:
        site_biomass, species, figure = biomass.fetch_biomass(
            affiliation, year, requested_site)
        
        doc.add_heading('Cover Crop Species and Biomass:', 4)

        if species is None:
            species = "Not available"
        elif len(species.index) > 0:
            tempstr = ""
            for index, row in species.iterrows():
                tempstr += row['cc_species'] + ", "
            species = tempstr[:-2]
        else:
            species = species['cc_species'].iloc[0]
       
        doc.add_paragraph('Cover Crop Species: ', style='List Bullet'
                          ).add_run(species)
        
        # Biomass and comparison
        if site_biomass is None:
            site_biomass = "Not available"
        elif len(site_biomass.index) > 0:
            tempstr = ""
            for index, row in site_biomass.iterrows():
                tempstr += str(round(row['uncorrected_cc_dry_biomass_lb_ac'])) + ", "
            site_biomass = tempstr[:-2]
        else:
            site_biomass = str(round(site_biomass['uncorrected_cc_dry_biomass_lb_ac'].iloc[0]))

        doc.add_paragraph('Dry Matter (lbs/acre):', style='List Bullet'
                          ).add_run(site_biomass)
        biomass_comp_para = doc.add_paragraph('Dry Matter in Comparison to Others in the Region:' +
                                              ' \n', style='List Bullet')
        try:
            figure.seek(0)
            doc.add_picture(figure)
        except Exception as no_graph:
            biomass_comp_para.add_run("Comparison not available")
        
    except Exception as e:
        print(e)
    finally:
        return doc


def doc_cropquality(doc, affilition, requested_site):
    try:
        nitrogen, carbohydrates, holo_cellulose, lignin = raw_nir.fetch_nir(
            affilition, requested_site)
        doc.add_heading('Cover Crop Quality:', 4)
        doc.add_paragraph('% Nitrogen: ', style='List Bullet'
                          ).add_run(str(round(nitrogen, 2)) if not np.isnan(nitrogen) else "Not available")
        doc.add_paragraph('% Carbohydrates: ', style='List Bullet'
                          ).add_run(str(round(carbohydrates, 2)) if not np.isnan(carbohydrates) else "Not available")
        doc.add_paragraph('% Holo-cellulose: ', style='List Bullet'
                          ).add_run(str(round(holo_cellulose, 2)) if not np.isnan(holo_cellulose) else "Not available")
        doc.add_paragraph('% Lignin: ', style='List Bullet'
                          ).add_run(str(round(lignin, 2)) if not np.isnan(nitrogen) else "Not available")

    except Exception as e:
        print(e)
    finally:
        return doc


def doc_yield(doc, affiliation, year, requested_site): 
    try:
        bare_yield, cover_yield, cash_crop, figure = crop_yield.fetch_yield(affiliation, year, requested_site)

        doc.add_heading('Yield:', 4)
        doc.add_paragraph('Cash Crop Species: ', style='List Bullet').add_run(cash_crop if cash_crop else "Not available")
        yield_para = doc.add_paragraph('Bare Ground: ', style='List Bullet')
        yield_para.add_run(
            str(round(bare_yield))+" bushels/ac" if bare_yield else "Not available")
        yield_para = doc.add_paragraph('Cover crop: ', style='List Bullet')
        yield_para.add_run(
            str(round(cover_yield))+" bushels/ac" if cover_yield else "Not available")
        
        try:
            figure.seek(0)
            doc.add_picture(figure)
        except Exception as no_graph:
            yield_para.add_run("Comparison not available")

    except Exception as e:
        print(e)
    finally:
        return doc


def doc_vwc(doc, requested_site, cash_planting, cash_harvest):
    try:
        if (cash_planting is not None):
            vwc, fig_temp, fig_moisture, fig_moisture_D, fig_moisture_C, fig_moisture_B, fig_moisture_A = moisture.fetch_vwc(cash_planting, cash_harvest, requested_site)
            
            vwc_dict = {}
            if vwc.size != 0:
                vwc['date'] = pd.to_datetime(vwc['timestamp'])
                new_df = (vwc.groupby(['treatment', pd.Grouper(
                    key='date', freq='W')]).agg({'vwc': 'mean'}))
                new_df = new_df.reset_index()
                new_df['date'] = pd.to_datetime(new_df['date']).dt.date
                new_df['date'] = pd.to_datetime(
                    new_df['date']).dt.strftime('%m/%d/%Y')

                for i in range(len(new_df)):
                    d = str(new_df.iloc[i].get('date'))
                    if d not in vwc_dict.keys():
                        vwc_dict[d] = ["", ""]
                    if str(new_df.iloc[i].get('treatment')) == "b":
                        vwc_dict[d][0] = str(round(new_df.iloc[i].get('vwc'), 1))
                    elif str(new_df.iloc[i].get('treatment')) == "c":
                        vwc_dict[d][1] = str(round(new_df.iloc[i].get('vwc'), 1))

            doc.add_heading('Soil Temperature: ', 4)
            doc.add_paragraph('Temperature Data: ', style='List Bullet')

            try:
                fig_temp.seek(0)
                doc.add_picture(fig_temp)
            except Exception as no_graph:
                doc.add_paragraph(
                    'Temperature Graph Not Available\n', style='List Bullet')

            #Water and Moisture
            doc.add_heading('Water and Soil Moisture: ', 4)
            doc.add_paragraph('Soil Moisture Data: ', style='List Bullet')
            doc.add_paragraph('Cover Crop vs Bare: ', style='List Number 2').add_run(
                "Refer the table below for overall(average) values")
            table = doc.add_table(rows=1, cols=3, style="Table Grid")
            row = table.rows[0].cells
            row[0].text = 'Week'
            row[1].text = 'Bare Ground (% volumetric water content)'
            row[2].text = 'Cover Crop (% volumetric water content)'

            for key, value in vwc_dict.items():
                # Adding a row and then adding data in it.
                row = table.add_row().cells
                # Converting id to string as table can only take string input
                row[0].text = key
                row[1].text = value[0]
                row[2].text = value[1]

            try:
                fig_moisture.seek(0)
                doc.add_picture(fig_moisture)
            except Exception as no_graph:
                doc.add_paragraph(
                    'Soil Moisture Graph not Available\n', style='List Bullet')
            try:
                fig_moisture_D.seek(0)
                doc.add_picture(fig_moisture_D)
            except Exception as no_graph:
                doc.add_paragraph(
                    'Soil Moisture D Graph not Available\n', style='List Bullet')
            try:
                fig_moisture_C.seek(0)
                doc.add_picture(fig_moisture_C)
            except Exception as no_graph:
                doc.add_paragraph(
                    'Soil Moisture C Graph not Available\n', style='List Bullet')
            try:
                fig_moisture_B.seek(0)
                doc.add_picture(fig_moisture_B)
            except Exception as no_graph:
                doc.add_paragraph(
                    'Soil Moisture B Graph not Available\n', style='List Bullet')
            try:
                fig_moisture_A.seek(0)
                doc.add_picture(fig_moisture_A)
            except Exception as no_graph:
                doc.add_paragraph(
                    'Soil Moisture A Graph not Available\n', style='List Bullet')

    except Exception as e:
        print(e)
    finally:
        return doc
    

def doc_dst_para(doc):
    try:
        doc.add_heading('Decision Support Tools:', 3)
        dst_para = doc.add_paragraph()
        dst_para.add_run('The Decision Support Tools (DSTs) are designed for farmers '+
                        'to input their data and receive custom generated information '+
                        'on how to address their management strategies.')
        
        dst_para = doc.add_paragraph(style='List Bullet')
        hyperlink.add_hyperlink(dst_para, 'Cover Crop Economic Decision Support Tool', 'https://covercrop-econ.org/')
        dst_para.add_run(': The Cover Crop Economic Decision Support Tool (DST) helps users '+
                         'better understand the impact of incorporating cover crop species or '+
                         'mixtures into their farm. Both the decision of what cover crop '+
                         'species to use and the resulting economic impact are very specific to '+
                         'individual operations. The DST will help you understand individualized '+
                         'conditions and management decisions.')
        
        dst_para = doc.add_paragraph(style='List Bullet')
        hyperlink.add_hyperlink(dst_para, 'Cover Crops Incentives Explorer Tool', 'https://covercrop-incentives.org/')
        dst_para.add_run(': Across the United States, there are many programs supporting the '+
                         'adoption of cover crops. This tool presents some of those programs, '+
                         'available at the federal and state levels, and summarizes their main '+
                         'characteristics.')

        dst_para = doc.add_paragraph(style='List Bullet')
        hyperlink.add_hyperlink(dst_para, 'Cover Crop Nitrogen Calculator', 'https://covercrop-ncalc.org/')
        dst_para.add_run(': Cover crops influence nitrogen (N) management to subsequent cash crops. '+
                         'Some of the N taken up or fixed by the cover crops becomes available over '+
                         'the cash crop growing season following termination. Estimating the rate of '+
                         'N release is challenging. The Cover Crop N Calculator provides a user-friendly '+
                         'approach to estimate decay of cover crop residues and release of N for offsetting '+
                         'N fertilizer inputs. This tool was developed for farmers and agricultural professionals.')

        dst_para = doc.add_paragraph(style='List Bullet')
        hyperlink.add_hyperlink(dst_para, 'Species Selector Tool', 'https://covercrop-selector.org/')
        dst_para.add_run(': The Cover Crop Species Selector is designed to help you choose cover crop '+
                         'species and explore their management characteristics. Data in this tool is '+
                         'curated from stakeholder experience by the four regional cover crop councils.')


    except Exception as e:
        print(e)
    finally:
        return doc