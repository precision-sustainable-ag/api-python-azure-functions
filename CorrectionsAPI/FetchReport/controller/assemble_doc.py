from ..controller import create_doc
from docx import Document
from datetime import datetime
import pandas as pd


def assemble(site_info, farm_hist, requested_site):
    cash_planting = farm_hist['cash_crop_planting_date'].iloc[0]
    if cash_planting == "" or cash_planting is None:
        cash_planting = None
    else:
        cash_planting = pd.to_datetime(cash_planting)

    cash_harvest = farm_hist['cash_crop_harvest_date'].iloc[0]
    if cash_harvest == "" or cash_harvest is None:
        cash_harvest = None
    else:
        cash_harvest = pd.to_datetime(cash_harvest)

    cover_planting = farm_hist['cc_planting_date'].iloc[0]
    if cover_planting == "" or cover_planting is None:
        cover_planting = None
    else:
        cover_planting = pd.to_datetime(cover_planting)

    cover_termination = farm_hist['cc_termination_date'].iloc[0]
    if cover_termination == "" or cover_termination is None:
        cover_termination = None
    else:
        cover_termination = pd.to_datetime(cover_termination)


    lat = round(site_info.iloc[0].get("latitude"), 4) \
        if site_info.iloc[0].get("latitude") is not None else None
    lon = round(site_info.iloc[0].get("longitude"), 4) \
        if site_info.iloc[0].get("latitude") is not None else None

    affilition = site_info.iloc[0].get("affiliation")

    doc = Document()
    doc = create_doc.doc_header(doc)

    # PSA paragraph
    doc.add_paragraph("\nThe Precision Sustainable Agriculture (PSA)" +
                      " On-Farm network deploys common research protocols that study the " +
                      "short-term effects of cover crops on farms that currently use cover" +
                      " crops. By utilizing farms with different management practices, the data" +
                      " collected can account for a wide range of factors such as termination " +
                      "timing, specific selections, and climate impacts. The data is used to " +
                      "build tools to aid in site-specific management decisions.\n")

    doc.add_heading('Farm Details', 2)
    doc = create_doc.doc_farmDetails(doc, site_info)

    # Dates this report summarizes
    doc.add_heading('Report Generated On:', 4)
    current_date = datetime.now()
    doc.add_paragraph(current_date.strftime("%m/%d/%Y"))

    # Summary of Activities and Management
    doc.add_heading('Summary of Activities and Management', 2)

    # Dates
    doc = create_doc.doc_cashcrop(doc, cash_planting, cash_harvest)
    doc = create_doc.doc_covercrop(doc, cover_planting, cover_termination)

    
    # GDD
    doc = create_doc.doc_gdd(
        doc, cash_planting, cash_harvest, cover_planting, cover_termination, lat, lon)

    # Precipitation
    doc = create_doc.doc_precipitation(
        doc, cash_planting, cash_harvest, cover_planting, cover_termination, lat, lon)

    # Biomass
    doc = create_doc.doc_biomass(doc, site_info, requested_site)

    # Composition
    doc = create_doc.doc_cropquality(doc, affilition, requested_site)

    # Yield
    doc = create_doc.doc_yield(doc, site_info, requested_site)

    # Temperature, Water and Moisture
    doc = create_doc.doc_vwc(doc, requested_site, cash_planting, cash_harvest)
    
    # Decision Support Tools
    doc = create_doc.doc_dst_para(doc)


    return doc
