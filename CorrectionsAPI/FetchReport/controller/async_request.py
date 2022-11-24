import asyncio
import io
import aiohttp
# from docx import Document
from docx.shared import Inches
import pandas as pd
import asyncio
from docx import Document
from ..controller import hyperlink
from datetime import datetime
import matplotlib.ticker as mticker
import pandas as pd
import matplotlib.pyplot as plt
plt.switch_backend('agg')

class async_request:
    url_biomass='https://api.precisionsustainableag.org/onfarm/biomass'
    url_yield='https://api.precisionsustainableag.org/onfarm/yield'
    url_gdd='https://api.precisionsustainableag.org/weather/daily'
    url_moisture='https://api.precisionsustainableag.org/onfarm/soil_moisture'
    url_precipitation='https://api.precisionsustainableag.org/weather/daily'
    url_nir='https://api.precisionsustainableag.org/onfarm/raw'

    api_header = { 'Accept': 'application/json',
                'x-api-key': 'e3d0dbf3-2884-43e3-9e5a-ad3d866df28c', }
    region = {
        "Northeast": ["PA", "VT", "NH"],
        "Southeast": ["AL", "FL", "GA", "NC"],
        "Mid-Atlantic": ["DE", "MD", "VA", "VAWest"],
        "Midwest": ["IN", "KS", "NE", "Indigo", "AR"],
    }
    aff_2_region = {
        "PA": "Northeast",
        "VT": "Northeast",
        "NH": "Northeast",
        "AL": "Southeast",
        "FL": "Southeast",
        "GA": "Southeast",
        "NC": "Southeast",
        "DE": "Mid-Atlantic",
        "MD": "Mid-Atlantic",
        "VA": "Mid-Atlantic",
        "VAWest": "Mid-Atlantic",
        "IN": "Midwest",
        "KS": "Midwest",
        "NE": "Midwest",
        "Indigo": "Midwest",
        "AR": "Midwest",
    }

    def __init__(self, doc, report_data, requested_site, cash_planting, \
        cash_harvest, cover_planting, cover_termination, lat, lon, affiliation):
        self.doc=doc
        self.report=report_data
        self.requested_site=requested_site
        self.cash_planting=cash_planting
        self.cash_harvest=cash_harvest
        self.cover_planting=cover_planting
        self.cover_termination=cover_termination
        self.lat=lat
        self.lon=lon
        self.requested_site=requested_site
        self.affiliation=affiliation
        affiliations = ",".join(self.region[self.aff_2_region[affiliation]])

        param_gdd1={'lat': lat, 'lon': lon, 'start': cash_planting, 'end': cash_harvest,\
                        'stats': 'sum(gdd)', 'gddbase': 10}
        param_gdd2={'lat': lat, 'lon': lon, 'start': cover_planting, 'end': cover_termination,\
                        'stats': 'sum(gdd)', 'gddbase': 4}
        param_precipitation1={'lat': lat, 'lon': lon, 'start': cash_planting,\
                        'end': cash_harvest, 'stats': 'sum(precipitation)'}
        param_precipitation2={'lat': lat, 'lon': lon, 'start': cover_planting,\
                        'end': cover_termination, 'stats': 'sum(precipitation)'}
        param_biomass={'affiliation': affiliations, 'output':'json'}
        param_nir={'table': 'biomass_nir', 'affiliation': affiliation, \
            'code': requested_site, 'output':'json'}
        param_yield={'code': requested_site}
        param_moisture={'type': 'tdr', 'start': cash_planting, 'code': requested_site}

        loop = asyncio.get_event_loop()
        urls = [(self.url_gdd, param_gdd1), (self.url_gdd, param_gdd2), \
            (self.url_precipitation, param_precipitation1), \
                (self.url_precipitation, param_precipitation2), \
                    (self.url_biomass, param_biomass), (self.url_nir, param_nir), \
                        (self.url_yield, param_yield), (self.url_moisture, param_moisture)]
        future = asyncio.ensure_future(self.fetch_all(urls))
        self.res_gdd1, self.res_gdd2, self.res_prec1, self.res_prec2, self.res_biomass, self.res_nir, self.res_yield, self.res_moisture = \
            loop.run_until_complete(future)

    async def fetch(self, session, url, params):
        async with session.get(url, params=params, headers=self.api_header) as response:
            return await response.json()

    async def fetch_all(self, urls):
        async with aiohttp.ClientSession() as session:
            res_gdd1, res_gdd2, res_prec1, res_prec2, res_biomass, res_nir, res_yield, res_moisture = \
                await asyncio.gather(*[self.fetch(session, url, params) for url, params in urls], return_exceptions=True)
            # print(results)
            return res_gdd1, res_gdd2, res_prec1, res_prec2, res_biomass, res_nir, res_yield, res_moisture

    def doc_header(self):
        try:
            # header section
            # adding header with PSA logo
            section = self.doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0]
            header_logo = header_para.add_run()
            header_logo.add_picture("FetchReport/PSA.png", width=Inches(1))

            # footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.add_run(
                'For questions about this report, ask: abc@xyz.com')
        except Exception as e:
            print(e)
        # finally:
        #     return doc

    def doc_main_para(self):
        # PSA paragraph
        self.doc.add_paragraph("\nThe Precision Sustainable Agriculture (PSA)" +
                        "On-Farm network deploys common research protocols that study the " +
                        "short-term effects of cover crops on farms that currently use cover" +
                        " crops. By utilizing farms with different management practices, the data" +
                        " collected can account for a wide range of factors such as termination " +
                        "timing, specific selections, and climate impacts. The data is used to " +
                        "build tools to aid in site-specific management decisions.\n")


    def doc_farmDetails(self):
        try:
            self.doc.add_heading('Farm Details', 2)
            # Farm Name
            self.doc.add_heading('Farm Name:', 4)
            self.doc.add_paragraph().add_run(
                self.report_data.iloc[0].get("code"))

            # Farm Address
            self.doc.add_heading('Farm Address:', 4)
            self.doc.add_paragraph().add_run(
                self.report_data.iloc[0].get("address"))

            lat = round(self.report_data.iloc[0].get("latitude"), 4)
            lon = round(self.report_data.iloc[0].get("longitude"), 4)

            # Site Description
            self.doc.add_heading('Site Description:', 4)
            gps_para = self.doc.add_paragraph()
            gps_para.add_run('GPS Co-ordinates\n').itallic = True
            gps_para.add_run('Latitude: ')
            gps_para.add_run(str(lat))
            gps_para.add_run('\t')
            gps_para.add_run('Longitude: ')
            gps_para.add_run(str(lon))
            gps_para.add_run('\n')

            maps_link = "https://www.google.com/maps/search/{lat},{lon}"\
                .format(lat=lat, lon=lon)
            hyperlink.add_hyperlink(gps_para, maps_link, maps_link)

            # Dates this report summarizes
            self.doc.add_heading('Dates this report summarizes:', 4)
            current_date = datetime.now()
            self.doc.add_paragraph(current_date.strftime("%m/%d/%Y"))

            # Summary of Activities and Management
            self.doc.add_heading('Summary of Activities and Management', 2)

        except Exception as e:
            print(e)
        # finally:
        #     return doc


    def doc_cashcrop(self):
        try:
            cash_days = (self.cash_harvest-self.cash_planting).days \
                if (self.cash_planting is not None and self.cash_harvest is not None) else None
            self.doc.add_heading('Cash crop days:', 4)
            # Cash crop dates
            self.doc.add_paragraph('Date of planting Cash crop: ', style='List Bullet'
                            ).add_run(self.cash_planting.strftime("%m/%d/%Y") if self.cash_planting else "Not yet entered")
            self.doc.add_paragraph('Date of harvest Cash crop: ', style='List Bullet'
                            ).add_run(self.cash_harvest.strftime("%m/%d/%Y") if self.cash_harvest else "Not yet entered")
            self.doc.add_paragraph('Cash crop no of days in production: ', style='List Bullet'
                            ).add_run(str(cash_days) if cash_days else "________")

        except Exception as e:
            print(e)
        # finally:
        #     return doc


    def doc_covercrop(self):
        try:
            cover_days = (self.cover_termination-self.cover_planting).days \
                if (self.cover_planting is not None and self.cover_termination is not None) else None
            self.doc.add_heading('Cover crop days:', 4)
            # Cover crop dates
            self.doc.add_paragraph('Date of planting Cover crop: ', style='List Bullet'
                            ).add_run(self.cover_planting.strftime("%m/%d/%Y") if self.cover_planting else "Not yet entered")
            self.doc.add_paragraph('Date of termination Cover crop: ', style='List Bullet'
                            ).add_run(self.cover_termination.strftime("%m/%d/%Y") if self.cover_termination else "Not yet entered")
            self.doc.add_paragraph('Cover crop no of days in production: ', style='List Bullet'
                            ).add_run(str(cover_days) if cover_days else "________")
        except Exception as e:
            print(e)
        # finally:
        #     return doc


    def doc_gdd(self):
        try:
            cash_gdd = self.res_gdd1 \
                if (self.cash_planting is not None and self.cash_harvest is not None) else None
            cover_gdd = self.res_gdd2 \
                if (self.cover_planting is not None and self.cover_termination is not None) else None

            self.doc.add_heading('GDD:', 4)
            gdd_para = self.doc.add_paragraph(
                'GDD for Cover crop is ', style='List Bullet')
            gdd_para.add_run(
                str(round(cover_gdd[0]['sum(gdd)'], 1)) if cover_gdd else "________")
            gdd_para = self.doc.add_paragraph(
                'GDD for Cash crop is ', style='List Bullet')
            gdd_para.add_run(
                str(round(cash_gdd[0]['sum(gdd)'], 1)) if cash_gdd else "________")

        except Exception as e:
            print(e)
        # finally:
        #     return doc


    def doc_precipitation(self):
        try:
            cash_precipitation = self.res_prec1 \
                if (self.cash_planting is not None and self.cash_harvest is not None) else None
            cover_precipitation = self.res_prec2 \
                if (self.cover_planting is not None and self.cover_termination is not None) else None

            self.doc.add_heading('Precipitation:', 4)
            # Precipitation
            preci_para = self.doc.add_paragraph(
                'Precipitation for Cover crop is ', style='List Bullet')
            preci_para.add_run(str(round(
                cover_precipitation[0]['sum(precipitation)'], 1))+" mm" if cover_precipitation else "________")
            preci_para = self.doc.add_paragraph(
                'Precipitation for Cash crop is ', style='List Bullet')
            preci_para.add_run(str(round(
                cash_precipitation[0]['sum(precipitation)'], 1))+" mm" if cash_precipitation else "________")

        except Exception as e:
            print(e)
        # finally:
        #     return doc

    def doc_biomass(self):
        try:
            biomass_data = pd.DataFrame(self.res_biomass)
            site_biomass = biomass_data[biomass_data['code'] == self.requested_site]
            if len(site_biomass) > 0 and \
                    pd.notna(site_biomass.iloc[0].get('ash_corrected_cc_dry_biomass_kg_ha')):
                biomass_data['ash_corrected_cc_dry_biomass_lb_ac'] = (
                    biomass_data['ash_corrected_cc_dry_biomass_kg_ha']*0.892179)
                site_year = site_biomass.iloc[0].get('year')
                biomass_data = biomass_data[biomass_data['year'] == site_year]
                biomass_data['Rank'] = (
                    biomass_data['ash_corrected_cc_dry_biomass_lb_ac'].rank(ascending=1))
                biomass_data.sort_values(
                    by=['ash_corrected_cc_dry_biomass_lb_ac'], inplace=True)
                site_biomass = (
                    biomass_data.loc[biomass_data['code'] == self.requested_site])

                yaxis = biomass_data['ash_corrected_cc_dry_biomass_lb_ac'].to_list(
                )
                xaxis = list(range(1, len(yaxis)+1))

                # create a figure and save the plot jpg image
                figure = io.BytesIO()
                plt.figure()
                plt.scatter(xaxis, yaxis, color='black', alpha=0.5)
                plt.scatter(int(site_biomass.iloc[0].get("Rank")), site_biomass.iloc[0].get(
                    "ash_corrected_cc_dry_biomass_lb_ac"), color='red', s=50)
                plt.text(int(site_biomass.iloc[0].get("Rank")), site_biomass.iloc[0].get(
                    "ash_corrected_cc_dry_biomass_lb_ac"), self.requested_site)
                # plt.title("Biomass data for {reg} region in year {year}".format(reg=aff_2_region[affiliation], year=str(site_year)))
                plt.title("This is your farm's dry matter in comparison to all farms that use \n cover crops in our network in the {reg} region".format(
                    reg=self.aff_2_region[self.affiliation]))
                plt.gca().xaxis.set_major_locator(mticker.MultipleLocator(2))
                plt.xlabel("Rank")
                plt.ylabel("Biomass produce in lbs/acre")
                plt.savefig(figure)
                # plt.savefig("FetchReport/Graph.png")
                plt.clf()
                plt.close()

                site_biomass, species, figure = site_biomass.iloc[0].get("ash_corrected_cc_dry_biomass_lb_ac"), \
                        site_biomass.iloc[0].get("cc_species"), figure
                self.doc.add_heading('Cover crop species and biomass:', 4)
                species = species if species else "Unavailable"
                self.doc.add_paragraph('Cover crop species: ', style='List Bullet'
                                ).add_run(species)
                # Biomass and comparison
                site_biomass = str(round(site_biomass, 1)
                                ) if site_biomass else "Unavailable"
                self.doc.add_paragraph('Dry matter (lbs/acre):', style='List Bullet'
                                ).add_run(site_biomass)
                biomass_comp_para = self.doc.add_paragraph('Dry matter in comparison to others in the region:' +
                                                    ' \n', style='List Bullet')
                try:
                    figure.seek(0)
                    self.doc.add_picture(figure)
                except Exception as no_graph:
                    biomass_comp_para.add_run("Comparison not available")

        except Exception as e:
            print(e)
        # finally:
        #     return doc


    def doc_cropquality(self):
        try:
            nir = pd.DataFrame(self.res_nir)
            nitrogen = None
            carbohydrates = None
            holo_cellulose = None
            lignin = None

            nitrogen = nir["percent_n_nir"].mean()
            carbohydrates = nir["percent_carb_nnorm"].mean()
            holo_cellulose = nir["percent_hemicell_calc"].mean(
            )+nir["percent_cellulose_calc"].mean()
            lignin = nir["percent_lignin_norm"].mean()

            self.doc.add_heading('Cover crop quality:', 4)
            self.doc.add_paragraph('% Nitrogen: ', style='List Bullet'
                            ).add_run(str(round(nitrogen, 2)))
            self.doc.add_paragraph('% Carbohydrates: ', style='List Bullet'
                            ).add_run(str(round(carbohydrates, 2)))
            self.doc.add_paragraph('% Holo-cellulose: ', style='List Bullet'
                            ).add_run(str(round(holo_cellulose, 2)))
            self.doc.add_paragraph('% Lignin: ', style='List Bullet'
                            ).add_run(str(round(lignin, 2)))

        except Exception as e:
            print(e)
        # finally:
        #     return doc


    def doc_yield(self):
        try:
            # Converting from Mg/ha to bushels/acre
            mg_to_bushels = {
                "Corn": 16,
                "Wheat": 15,
                "Soybeans": 15,
                "Cotton": 28
            }
            yield_data = pd.DataFrame(self.res_yield)
            bare_yield = None
            cover_yield = None
            if len(yield_data) > 0 and 'adjusted.grain.yield.Mg_ha' in yield_data.columns:
                if len(yield_data[yield_data['treatment'] == 'B']) > 0:
                    bare_yield_data = yield_data[yield_data['treatment'] == 'B']
                    if bare_yield_data.iloc[0]['adjusted.grain.yield.Mg_ha']:
                        bare_yield = bare_yield_data.iloc[0]['adjusted.grain.yield.Mg_ha']*\
                            mg_to_bushels[bare_yield_data.iloc[0]['cash.crop']]
                if len(yield_data[yield_data['treatment'] == 'C']) > 0:
                    cover_yield_data = yield_data[yield_data['treatment'] == 'C']
                    if cover_yield_data.iloc[0]['adjusted.grain.yield.Mg_ha']:
                        cover_yield = cover_yield_data.iloc[0]['adjusted.grain.yield.Mg_ha']*\
                            mg_to_bushels[cover_yield_data.iloc[0]['cash.crop']]
            self.doc.add_heading('Yield:', 4)
            yield_para = self.doc.add_paragraph('Bare Ground: ', style='List Bullet')
            yield_para.add_run(
                str(round(bare_yield))+" bushels/ac" if bare_yield else "Not available")
            yield_para = self.doc.add_paragraph('Cover: ', style='List Bullet')
            yield_para.add_run(
                str(round(cover_yield))+" bushels/ac" if cover_yield else "Not available")

        except Exception as e:
            print(e)
        # finally:
        #     return doc

    def plot_graph(vwc, file_name, start_date, end_date, depth="overall"):
        try:
            fig = io.BytesIO()
            vwc['date'] = pd.to_datetime(vwc['timestamp'])
            new_df = (vwc.groupby(['treatment', pd.Grouper(
                key='date', freq='W')]).agg({'vwc': 'mean'}))
            new_df = new_df.reset_index()
            new_df['date'] = pd.to_datetime(new_df['date']).dt.date
            bare_df = new_df.loc[new_df['treatment'] == 'b']
            cover_df = new_df.loc[new_df['treatment'] == 'c']
            bare_df.sort_values(by=['date'], inplace=True)
            cover_df.sort_values(by=['date'], inplace=True)

        # new_df['date'] = pd.to_datetime(new_df['date']).dt.strftime('%m/%d/%Y')
            y_bare = bare_df['vwc'].to_list()
            y_cover = cover_df['vwc'].to_list()
            x_bare = bare_df['date'].to_list()
            x_cover = cover_df['date'].to_list()

            f = plt.figure()
            f.set_figwidth(6)
            f.set_figheight(3.75)
            # Plotting both the curves simultaneously
            plt.plot(x_bare, y_bare, color='y', label='Bare Ground')
            plt.plot(x_cover, y_cover, color='g', label='Cover crop')
            plt.gcf().autofmt_xdate()

            # Naming the x-axis, y-axis and the whole graph
            plt.xlabel("Week")
            plt.ylabel("Moisture(%)")
            plt.title("Soil Moisture percentage at {depth} depth"
                    .format(depth=depth))

            # Adding legend, which helps us recognize curve according to it's color
            plt.legend()
            plt.savefig(fig)
            # plt.savefig(file_name)
            plt.clf()
            plt.close()
            return fig
        except Exception as e:
            print(e)


    def doc_vwc(self):
        try:
            start_date = self.cash_planting
            end_date = self.cash_harvest
            vwc_data = pd.DataFrame(self.res_moisture)
            if len(vwc_data) != 0:
                vwc_overall = vwc_data

                fig_moisture = self.plot_graph(vwc_overall, "FetchReport/data/MoistureGraph.png",
                                        start_date, end_date)
                vwc_d = vwc_data[vwc_data["center_depth"] == -5]
                fig_moisture_D = self.plot_graph(vwc_d, "FetchReport/data/MoistureGraphD.png",
                                            start_date, end_date, "surface")
                vwc_c = vwc_data[vwc_data["center_depth"] == -15]
                fig_moisture_C = self.plot_graph(vwc_c, "FetchReport/data/MoistureGraphC.png",
                                            start_date, end_date, "6 inch")
                vwc_b = vwc_data[vwc_data["center_depth"] == -45]
                fig_moisture_B = self.plot_graph(vwc_b, "FetchReport/data/MoistureGraphB.png",
                                            start_date, end_date, "18 inch")
                vwc_a = vwc_data[vwc_data["center_depth"] == -80]
                fig_moisture_A = self.plot_graph(vwc_a, "FetchReport/data/MoistureGraphA.png",
                                            start_date, end_date, "31 inch")

                # vwc_d['date'] = pd.to_datetime(vwc_d['timestamp'])
                new_df = (vwc_d.groupby(['treatment', pd.Grouper(
                    key='date', freq='W')]).agg({'soil_temp': 'mean'}))
                new_df = new_df.reset_index()
                new_df['date'] = pd.to_datetime(new_df['date']).dt.date
                bare_df = new_df.loc[new_df['treatment'] == 'b']
                cover_df = new_df.loc[new_df['treatment'] == 'c']
                bare_df.sort_values(by=['date'], inplace=True)
                cover_df.sort_values(by=['date'], inplace=True)

                # new_df['date']=pd.to_datetime(new_df['date']).dt.strftime('%m/%d/%Y')
                bare_df['soil_temp'] = bare_df['soil_temp'].apply(
                    lambda x: x*1.8 + 32)
                cover_df['soil_temp'] = cover_df['soil_temp'].apply(
                    lambda x: x*1.8 + 32)
                y_bare = bare_df['soil_temp'].to_list()
                y_cover = cover_df['soil_temp'].to_list()
                x_bare = bare_df['date'].to_list()
                x_cover = cover_df['date'].to_list()

                # Plotting both the curves simultaneously
                plt.plot(x_bare, y_bare, color='y', label='Bare Ground')
                plt.plot(x_cover, y_cover, color='g', label='Cover crop')
                plt.gcf().autofmt_xdate()

                # Naming the x-axis, y-axis and the whole graph
                plt.xlabel("Week")
                plt.ylabel("Temperature in deg F")
                plt.title("Soil temperature at surface")

                plt.legend()
                fig_temp = io.BytesIO()
                plt.savefig(fig_temp)
                # plt.savefig("FetchReport/data/TemperatureGraph.png")
                plt.clf()
                plt.close()
                # To load the display window

            vwc = vwc_data\
                if (self.cash_planting is not None) else None
            # if (cash_planting is not None and cash_harvest is not None) else None
            vwc_dict = {}
            if vwc.size != 0:
                vwc['date'] = pd.to_datetime(vwc['timestamp'])
                new_df = (vwc.groupby(['treatment', pd.Grouper(
                    key='date', freq='W')]).agg({'vwc': 'mean'}))
                new_df = new_df.reset_index()
                new_df['date'] = pd.to_datetime(new_df['date']).dt.date
                new_df['date'] = pd.to_datetime(
                    new_df['date']).dt.strftime('%m/%d/%Y')

                for i in range(len(new_df)):
                    d = str(new_df.iloc[i].get('date'))
                    if d not in vwc_dict.keys():
                        vwc_dict[d] = ["", ""]
                    if str(new_df.iloc[i].get('treatment')) == "b":
                        vwc_dict[d][0] = str(round(new_df.iloc[i].get('vwc'), 3))
                    elif str(new_df.iloc[i].get('treatment')) == "c":
                        vwc_dict[d][1] = str(round(new_df.iloc[i].get('vwc'), 3))

            self.doc.add_heading('Soil Temperature: ', 4)
            self.doc.add_paragraph('Temperature data: ', style='List Bullet')

            try:
                fig_temp.seek(0)
                self.doc.add_picture(fig_temp)
            except Exception as no_graph:
                self.doc.add_paragraph(
                    'Temperature Graph not available\n', style='List Bullet')

            #Water and Moisture
            self.doc.add_heading('Water and Moisture: ', 4)
            self.doc.add_paragraph('Moisture data: ', style='List Bullet')
            self.doc.add_paragraph('Cover crop vs bare: ', style='List Number 2').add_run(
                "Refer the table below for overall(avr) values")
            table = self.doc.add_table(rows=1, cols=3, style="Table Grid")
            row = table.rows[0].cells
            row[0].text = 'Week'
            row[1].text = 'Bare Ground'
            row[2].text = 'Cover Crop'

            for key, value in vwc_dict.items():
                # Adding a row and then adding data in it.
                row = table.add_row().cells
                # Converting id to string as table can only take string input
                row[0].text = key
                row[1].text = value[0]
                row[2].text = value[1]

            try:
                fig_moisture.seek(0)
                self.doc.add_picture(fig_moisture)
            except Exception as no_graph:
                self.doc.add_paragraph(
                    'Moisture Graph not available\n', style='List Bullet')
            try:
                fig_moisture_D.seek(0)
                self.doc.add_picture(fig_moisture_D)
            except Exception as no_graph:
                self.doc.add_paragraph(
                    'Moisture D Graph not available\n', style='List Bullet')
            try:
                fig_moisture_C.seek(0)
                self.doc.add_picture(fig_moisture_C)
            except Exception as no_graph:
                self.doc.add_paragraph(
                    'Moisture C Graph not available\n', style='List Bullet')
            try:
                fig_moisture_B.seek(0)
                self.doc.add_picture(fig_moisture_B)
            except Exception as no_graph:
                self.doc.add_paragraph(
                    'Moisture B Graph not available\n', style='List Bullet')
            try:
                fig_moisture_A.seek(0)
                self.doc.add_picture(fig_moisture_A)
            except Exception as no_graph:
                self.doc.add_paragraph(
                    'Moisture A Graph not available\n', style='List Bullet')

        except Exception as e:
            print(e)
        # finally:
        #     return doc

    def doc_end_para(self):
        # Decision Support Tools
        self.doc.add_heading('Decision Support Tools:', 3)
        dst_para = self.doc.add_paragraph()
        dst_para.add_run('The Decision Support Tools (DSTs) are designed for farmers' +
                        ' to input their data and receive custom generated information on how to ' +
                        'address their management strategies. The Cover Crop Nitrogen Calculator ' +
                        '(CC-NCALC) calculates the amount of nitrogen available after the planting ' +
                        'and termination of a cover crop. ')
    def conclude(self):
        return self.doc